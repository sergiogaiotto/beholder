"""Inventário do BEHOLDER_DATA_DIR — Pré-A do plano Empreiteiras-WF.

Para cada arquivo bruto (XLSX SAP, PDFs zipados, MSRV5 LPU, Analítico WF, DOCX),
extrai metadata estrutural (sem ler dados — só schemas + amostras) e gera:

- docs/DATA_INVENTORY.json  — manifesto completo, machine-readable
- docs/DATA_INVENTORY.md    — relatório humano para revisão

Roda com a venv do projeto:
    .venv\\Scripts\\python.exe scripts\\inventory_data.py
"""

from __future__ import annotations

import hashlib
import io
import json
import os
import sys
import zipfile
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import openpyxl  # type: ignore[import-untyped]
from docx import Document  # type: ignore[import-untyped]

DATA_DIR_ENV = "BEHOLDER_DATA_DIR"
DEFAULT_DATA_DIR = r"C:\_PERSONAL\beholder_data"

REPO_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = REPO_ROOT / "docs"
JSON_OUT = DOCS_DIR / "DATA_INVENTORY.json"
MD_OUT = DOCS_DIR / "DATA_INVENTORY.md"

CHUNK = 1024 * 1024  # 1 MB
TXT_SAMPLE_HEAD = 30
TXT_SAMPLE_TAIL = 10
DOCX_PREVIEW_CHARS = 800


@dataclass
class FileRecord:
    path: str
    relative_path: str
    size_bytes: int
    sha256: str
    kind: str  # xlsx | zip | txt | docx | other
    details: dict[str, Any] = field(default_factory=dict)


def sha256_of(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        while chunk := f.read(CHUNK):
            h.update(chunk)
    return h.hexdigest()


def inspect_xlsx(path: Path) -> dict[str, Any]:
    """Lê headers e dimensões de cada sheet sem carregar dados."""
    out: dict[str, Any] = {"sheets": []}
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # max_row/max_column vêm do dimension tag; podem ser unreliable em
            # arquivos exportados por SAP. Quando >0, confiamos; senão, sondamos.
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0
            header: list[str] = []
            second_row: list[str] = []
            row_iter = ws.iter_rows(values_only=True)
            try:
                first = next(row_iter)
                header = [str(c) if c is not None else "" for c in first]
                second = next(row_iter, None)
                if second is not None:
                    second_row = [str(c) if c is not None else "" for c in second]
            except StopIteration:
                pass
            out["sheets"].append(
                {
                    "name": sheet_name,
                    "max_row_declared": max_row,
                    "max_col_declared": max_col,
                    "header": header,
                    "header_count": len(header),
                    "second_row_sample": second_row,
                }
            )
    finally:
        wb.close()
    return out


def inspect_zip(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as z:
        entries = []
        for info in z.infolist():
            entries.append(
                {
                    "name": info.filename,
                    "size": info.file_size,
                    "compressed_size": info.compress_size,
                    "is_pdf": info.filename.lower().endswith(".pdf"),
                }
            )
    pdfs = [e for e in entries if e["is_pdf"]]
    return {
        "entries": entries,
        "entry_count": len(entries),
        "pdf_count": len(pdfs),
        "total_pdf_bytes": sum(e["size"] for e in pdfs),
    }


def inspect_txt(path: Path) -> dict[str, Any]:
    """Conta linhas (streaming) + amostra cabeça e cauda."""
    line_count = 0
    head: list[str] = []
    tail: list[str] = []
    # Streaming line count + collect head
    with path.open("r", encoding="utf-8", errors="replace") as f:
        for line in f:
            line_count += 1
            if len(head) < TXT_SAMPLE_HEAD:
                head.append(line.rstrip("\n"))
    # Tail: re-open and seek near end for last N lines
    if line_count > TXT_SAMPLE_HEAD:
        with path.open("rb") as f:
            f.seek(0, os.SEEK_END)
            end = f.tell()
            block = 64 * 1024
            data = b""
            while f.tell() > 0 and data.count(b"\n") <= TXT_SAMPLE_TAIL + 2:
                step = min(block, f.tell())
                f.seek(-step, os.SEEK_CUR)
                data = f.read(step) + data
                f.seek(-step, os.SEEK_CUR)
            tail = data.decode("utf-8", errors="replace").splitlines()[-TXT_SAMPLE_TAIL:]
    return {
        "line_count": line_count,
        "head_lines": head,
        "tail_lines": tail,
    }


def inspect_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    table_count = len(doc.tables)
    table_summary = []
    for i, tbl in enumerate(doc.tables[:10]):
        rows = len(tbl.rows)
        cols = len(tbl.rows[0].cells) if rows > 0 else 0
        table_summary.append({"index": i, "rows": rows, "cols": cols})
    return {
        "paragraph_count": len(paragraphs),
        "total_chars": len(full_text),
        "preview": full_text[:DOCX_PREVIEW_CHARS],
        "table_count": table_count,
        "table_summary": table_summary,
    }


def walk(root: Path) -> list[FileRecord]:
    records: list[FileRecord] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(root).as_posix()
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            kind = "xlsx"
            details = inspect_xlsx(path)
        elif suffix == ".zip":
            kind = "zip"
            details = inspect_zip(path)
        elif suffix == ".txt":
            kind = "txt"
            details = inspect_txt(path)
        elif suffix == ".docx":
            kind = "docx"
            details = inspect_docx(path)
        else:
            kind = "other"
            details = {}
        print(f"  inspected: {rel}", file=sys.stderr, flush=True)
        records.append(
            FileRecord(
                path=str(path),
                relative_path=rel,
                size_bytes=path.stat().st_size,
                sha256=sha256_of(path),
                kind=kind,
                details=details,
            )
        )
    return records


def fmt_size(n: int) -> str:
    if n < 1024:
        return f"{n} B"
    if n < 1024 ** 2:
        return f"{n / 1024:.1f} KB"
    if n < 1024 ** 3:
        return f"{n / 1024 ** 2:.1f} MB"
    return f"{n / 1024 ** 3:.2f} GB"


def render_markdown(records: list[FileRecord], data_dir: Path) -> str:
    total = sum(r.size_bytes for r in records)
    by_kind: dict[str, list[FileRecord]] = {}
    for r in records:
        by_kind.setdefault(r.kind, []).append(r)

    lines: list[str] = []
    lines.append("# DATA_INVENTORY — Empreiteiras-WF (Pré-A)")
    lines.append("")
    lines.append(
        f"Gerado em {datetime.now(timezone.utc).isoformat(timespec='seconds')} a partir de `{data_dir}`."
    )
    lines.append(
        "Fonte da verdade: `docs/DATA_INVENTORY.json`. Este MD é o resumo navegável."
    )
    lines.append("")
    lines.append(f"**Total**: {len(records)} arquivos, {fmt_size(total)}.")
    lines.append("")
    lines.append("## Resumo por tipo")
    lines.append("")
    lines.append("| Tipo | Qtd | Tamanho total |")
    lines.append("|---|---:|---:|")
    for kind in sorted(by_kind):
        items = by_kind[kind]
        lines.append(
            f"| `{kind}` | {len(items)} | {fmt_size(sum(r.size_bytes for r in items))} |"
        )
    lines.append("")

    # XLSX detail
    if by_kind.get("xlsx"):
        lines.append("## XLSX — schemas por sheet")
        lines.append("")
        for r in by_kind["xlsx"]:
            lines.append(f"### `{r.relative_path}`")
            lines.append("")
            lines.append(
                f"- size: **{fmt_size(r.size_bytes)}** · sha256: `{r.sha256[:16]}…`"
            )
            for sh in r.details.get("sheets", []):
                lines.append(
                    f"- sheet **{sh['name']}** — declared rows={sh['max_row_declared']:,}, cols={sh['max_col_declared']:,}, header_count={sh['header_count']}"
                )
                hdr_preview = sh["header"][:15]
                more = "" if len(sh["header"]) <= 15 else f" … (+{len(sh['header'])-15} cols)"
                if hdr_preview:
                    lines.append(f"  - header: `{hdr_preview}`{more}")
            lines.append("")

    # ZIP detail
    if by_kind.get("zip"):
        lines.append("## ZIP — contratos PDF")
        lines.append("")
        lines.append("| ZIP | empreiteira (path) | PDFs | bytes pdf |")
        lines.append("|---|---|---:|---:|")
        for r in sorted(by_kind["zip"], key=lambda x: x.relative_path):
            parts = r.relative_path.split("/")
            emp = parts[-2] if len(parts) >= 2 else "(root)"
            lines.append(
                f"| `{parts[-1]}` | {emp} | {r.details['pdf_count']} | {fmt_size(r.details['total_pdf_bytes'])} |"
            )
        lines.append("")
        total_pdfs = sum(r.details["pdf_count"] for r in by_kind["zip"])
        total_pdf_bytes = sum(r.details["total_pdf_bytes"] for r in by_kind["zip"])
        lines.append(
            f"**Total**: {len(by_kind['zip'])} ZIPs, {total_pdfs} PDFs, {fmt_size(total_pdf_bytes)} descomprimido."
        )
        lines.append("")

    # TXT detail
    if by_kind.get("txt"):
        lines.append("## TXT — relatórios SAP")
        lines.append("")
        for r in by_kind["txt"]:
            lines.append(f"### `{r.relative_path}`")
            lines.append(
                f"- size: **{fmt_size(r.size_bytes)}** · linhas: **{r.details['line_count']:,}** · sha256: `{r.sha256[:16]}…`"
            )
            lines.append("- head:")
            lines.append("```")
            for ln in r.details["head_lines"][:10]:
                lines.append(ln)
            lines.append("```")
            lines.append("")

    # DOCX detail
    if by_kind.get("docx"):
        lines.append("## DOCX")
        lines.append("")
        for r in by_kind["docx"]:
            lines.append(f"### `{r.relative_path}`")
            lines.append(
                f"- size: **{fmt_size(r.size_bytes)}** · parágrafos: {r.details['paragraph_count']} · tabelas: {r.details['table_count']}"
            )
            preview = r.details["preview"].replace("\n", "\n> ")
            lines.append("- preview:")
            lines.append("> " + preview)
            lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## Hash manifest (auditoria)")
    lines.append("")
    lines.append("| Arquivo | SHA-256 | Tamanho |")
    lines.append("|---|---|---:|")
    for r in sorted(records, key=lambda x: x.relative_path):
        lines.append(f"| `{r.relative_path}` | `{r.sha256}` | {fmt_size(r.size_bytes)} |")
    lines.append("")
    return "\n".join(lines)


def main() -> int:
    data_dir = Path(os.environ.get(DATA_DIR_ENV, DEFAULT_DATA_DIR))
    if not data_dir.exists():
        print(f"ERRO: {data_dir} não existe. Defina {DATA_DIR_ENV}.", file=sys.stderr)
        return 1
    print(f"Inspecionando {data_dir} …", file=sys.stderr, flush=True)
    records = walk(data_dir)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    JSON_OUT.write_text(
        json.dumps(
            {
                "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
                "data_dir": str(data_dir),
                "file_count": len(records),
                "total_bytes": sum(r.size_bytes for r in records),
                "files": [asdict(r) for r in records],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    MD_OUT.write_text(render_markdown(records, data_dir), encoding="utf-8")

    print(f"\nOK — gravado:", file=sys.stderr)
    print(f"  {JSON_OUT}", file=sys.stderr)
    print(f"  {MD_OUT}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
