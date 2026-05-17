"""Extrai conteúdo completo do DOCX para análise das regras originais."""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document  # type: ignore[import-untyped]

DOC = Path(r"C:\_PERSONAL\beholder_data\Regras - POC - Automation AnyWhere - IA na Monitoria de Pagamentos.docx")

def main() -> int:
    if not DOC.exists():
        print(f"Não encontrado: {DOC}", file=sys.stderr)
        return 1
    doc = Document(str(DOC))
    out: list[str] = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if not text:
            continue
        prefix = f"[{style}] " if style and style != "Normal" else ""
        out.append(f"{prefix}{text}")
    # Tabelas (se houver)
    for i, tbl in enumerate(doc.tables):
        out.append(f"\n--- TABELA {i} ({len(tbl.rows)} linhas) ---")
        for row in tbl.rows:
            out.append(" | ".join(cell.text.strip() for cell in row.cells))
    print("\n".join(out))
    return 0

if __name__ == "__main__":
    sys.exit(main())
